#!/usr/bin/env python3
"""
THỢ DEEPSEEK — Cập nhật giá sản phẩm từ file DOCX chuẩn
Yêu cầu: đọc file DOCX trong /media/huu-minh/DATA/tai-lieu/bảng giá các hãng/báo giá chuẩn/
và cập nhật js/prices.js với giá chuẩn từ file.

CẤM: thêm brand ngoài file chuẩn, tự ý tạo số liệu.
CHỈ: dùng dữ liệu từ DOCX.

Kết quả đầu ra: file js/prices.js đã cập nhật.
"""
import docx, re, os, json, sys

BASE = "/media/huu-minh/DATA/tai-lieu/bảng giá các hãng/báo giá chuẩn"
WEB_DIR = "/home/huu-minh/.openclaw/workspace/web-vlht"

def clean_price(s):
    s = s.replace('.','').replace(',','').replace('đ','').replace(' ','').replace('VNĐ','').replace('₫','').strip()
    m = re.search(r'\d+', s)
    return int(m.group(0)) if m else None

def fmt_price(n):
    if n is None: return "'Liên hệ'"
    s = f"{n:,}đ".replace(',', '.')
    return f"'{s}'"

def parse_munich():
    doc = docx.Document(os.path.join(BASE, "Bang-gia-Munich-TranHuuMinh.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            code = cells[0]
            name = cells[1]
            if code == name or not code: continue
            spec = cells[2]
            price_raw = cells[3]
            p = clean_price(price_raw.split('/')[0].strip())
            if p:
                prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

def parse_dulux():
    doc = docx.Document(os.path.join(BASE, "bang-gia-dulux-maxilite-22-05-2026.docx"))
    dulux, maxilite = [], []
    for ti in range(2):
        table = doc.tables[ti]
        is_max = (ti == 1)
        target = maxilite if is_max else dulux
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            name = cells[0]
            if len(set(cells[:3])) == 1: continue
            code = cells[1]
            spec = cells[2]
            price_raw = cells[3]
            p = clean_price(price_raw)
            if p:
                target.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return dulux, maxilite

def parse_jotun():
    doc = docx.Document(os.path.join(BASE, "bang-gia-jotun-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            name = cells[0]
            if 'SẢN PHẨM' in name: continue
            if len(set(cells[2:4])) == 1 and cells[2] != '': continue
            spec = cells[2]
            price_raw = cells[3]
            p = clean_price(price_raw)
            if p:
                prods.append({'code': f"JT_{name[:8]}", 'name': name, 'spec': spec, 'price': p})
    return prods

def parse_kova():
    doc = docx.Document(os.path.join(BASE, "bang-gia-kova-22-05-2026.docx"))
    prods = []
    for ti in range(3):
        table = doc.tables[ti]
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            code = cells[0]
            name = cells[1]
            if not code or not name: continue
            if code == name or 'SƠN NƯỚC' in code or 'MATÍT' in code: continue
            if code.startswith('(như'): continue
            spec = cells[2]
            price_raw = cells[3]
            p = clean_price(price_raw)
            if p:
                prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

def parse_nano():
    doc = docx.Document(os.path.join(BASE, "bang-gia-nano-house-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 3: continue
            name = cells[0]
            if 'SẢN PHẨM' in name: continue
            if len(set(cells[2:4])) == 1 and cells[2] != '': continue
            spec = cells[2]
            price_raw = cells[3] if len(cells) > 3 else cells[2]
            p = clean_price(price_raw)
            if p:
                code_match = re.search(r'[-–]\s*(\w+)', name)
                code = code_match.group(1) if code_match else name[:8]
                prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

def parse_nippon():
    doc = docx.Document(os.path.join(BASE, "bang-gia-nippon-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 3: continue
            name = cells[0]
            if 'SẢN PHẨM' in name: continue
            if len(set(cells[2:4])) == 1 and cells[2] != '': continue
            spec = cells[2]
            price_raw = cells[3] if len(cells) > 3 else ''
            p = clean_price(price_raw)
            if p:
                prods.append({'code': f"NP_{name[:10]}", 'name': name, 'spec': spec, 'price': p})
    return prods

def parse_sika():
    doc = docx.Document(os.path.join(BASE, "bang-gia-sika-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            name = cells[0]
            if name.startswith('I.') or name.startswith('II.') or name.startswith('III.'): continue
            if name.startswith('1.') or name.startswith('2.') or name.startswith('3.'): continue
            if len(set(cells[2:5])) == 1: continue
            spec = cells[2]
            price_raw = cells[4] if len(cells) > 4 else cells[3]
            p = clean_price(price_raw)
            if p:
                prods.append({'code': name[:15], 'name': name, 'spec': spec, 'price': p})
    return prods

def generate_prices_js(all_prices):
    """Generate the PRICES object for prices.js"""
    lines = []
    lines.append("// BẢNG GIÁ SẢN PHẨM - TỪ FILE CHUẨN ĐÃ DUYỆT (05/2026)")
    lines.append("// Nguồn: /media/.../báo giá chuẩn/")
    lines.append("")
    lines.append("const PRICES = {")
    
    for brand_id, prods in all_prices.items():
        if not prods: continue
        lines.append(f"  {brand_id}: {{")
        for p in prods:
            spec_clean = p['spec'].replace("'", "\\'").replace('"', '\\"')
            price_str = fmt_price(p['price'])
            lines.append(f"    '{p['code']}': {{ price: {price_str}, spec: '{spec_clean}' }},")
        # Remove trailing comma from last item (fine for JS)
        lines.append("  },")
    
    lines.append("};")
    lines.append("")
    lines.append("function getPrice(brand, code) {")
    lines.append("  const b = PRICES[brand];")
    lines.append("  if (!b || !b[code]) return null;")
    lines.append("  return b[code];")
    lines.append("}")
    
    return "\n".join(lines)

if __name__ == "__main__":
    print("=== THỢ DEEPSEEK: Cập nhật giá từ file chuẩn ===")
    
    all_prices = {}
    all_prices['munich'] = parse_munich()
    d, m = parse_dulux()
    all_prices['dulux'] = d
    all_prices['maxilite'] = m
    all_prices['jotun'] = parse_jotun()
    all_prices['kova'] = parse_kova()
    all_prices['nano'] = parse_nano()
    all_prices['nippon'] = parse_nippon()
    all_prices['sika'] = parse_sika()
    
    for k, v in all_prices.items():
        print(f"  {k}: {len(v)} products")
    
    # Cập nhật prices.js
    js_content = generate_prices_js(all_prices)
    output_path = os.path.join(WEB_DIR, "js", "prices.js")
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(js_content)
    
    print(f"\n✅ Đã cập nhật {output_path}")
    print(f"   {sum(len(v) for v in all_prices.values())} tổng sản phẩm có giá")
