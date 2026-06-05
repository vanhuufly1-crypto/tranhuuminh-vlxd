#!/usr/bin/env python3
"""Generate products.js + prices.js from approved DOCX files"""
import docx, re, os, json

BASE = "/media/huu-minh/DATA/tai-lieu/bảng giá các hãng/báo giá chuẩn"

def clean_price(s):
    s = s.replace('.','').replace(',','').replace('đ','').replace(' ','').replace('VNĐ','').replace('₫','').strip()
    m = re.search(r'\d+', s)
    return int(m.group(0)) if m else None

def fmt_price(n):
    if n is None: return "'Liên hệ'"
    s = f"{n:,}đ"
    return s

# ======== MUNICH ========
def parse_munich():
    doc = docx.Document(os.path.join(BASE, "Bang-gia-Munich-TranHuuMinh.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            code = cells[0]
            name = cells[1]
            if code == name or not code: continue
            spec = cells[2]
            price_raw = cells[3]
            first_price = price_raw.split('/')[0].strip()
            p = clean_price(first_price)
            prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    # Deduplicate by code (keep first occurrence)
    seen = set()
    unique = []
    for p in prods:
        if p['code'] not in seen:
            seen.add(p['code'])
            unique.append(p)
    return unique

# ======== DULUX + MAXILITE ========
def parse_dulux():
    doc = docx.Document(os.path.join(BASE, "bang-gia-dulux-maxilite-22-05-2026.docx"))
    dulux_list, maxilite_list = [], []
    
    for ti in range(2):
        table = doc.tables[ti]
        is_maxilite = (ti == 1)
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            name = cells[0]
            if len(set(cells[:3])) == 1: continue
            code = cells[1]
            spec = cells[2]
            price_raw = cells[3]
            p = clean_price(price_raw)
            item = {'code': code, 'name': name, 'spec': spec, 'price': p}
            unique_key = f"{code}_{name}"
            if is_maxilite:
                if not any(x['code']==code and x['spec']==spec for x in maxilite_list):
                    maxilite_list.append(item)
            else:
                if not any(x['code']==code and x['spec']==spec for x in dulux_list):
                    dulux_list.append(item)
    return dulux_list, maxilite_list

# ======== JOTUN ========
def parse_jotun():
    doc = docx.Document(os.path.join(BASE, "bang-gia-jotun-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            name = cells[0]
            if 'SẢN PHẨM' in name: continue
            if len(set(cells[2:4])) == 1 and cells[2] != '': continue
            spec = cells[2]
            price_raw = cells[3]
            p = clean_price(price_raw)
            if p:
                prods.append({'code': f"JT-{name[:4].upper()}", 'name': name, 'spec': spec, 'price': p})
    return prods

# ======== KOVA ========
def parse_kova():
    doc = docx.Document(os.path.join(BASE, "bang-gia-kova-22-05-2026.docx"))
    prods = []
    for ti in range(3):
        table = doc.tables[ti]
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            code = cells[0]
            name = cells[1]
            if not code or not name: continue
            if code == name or 'SƠN NƯỚC' in code or 'MATÍT' in code: continue
            if code.startswith('(như'): continue
            spec = cells[2]
            price_raw = cells[3]
            p = clean_price(price_raw)
            if p and not any(x['code']==code and x['spec']==spec for x in prods):
                prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

# ======== NANO ========
def parse_nano():
    doc = docx.Document(os.path.join(BASE, "bang-gia-nano-house-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 3: continue
            name = cells[0]
            if 'SẢN PHẨM' in name: continue
            if len(set(cells[2:4])) == 1 and cells[2] != '': continue
            spec = cells[2]
            price_raw = cells[3] if len(cells) > 3 else cells[2]
            p = clean_price(price_raw)
            if p:
                code_match = re.search(r'[-–]\s*(\w+)', name)
                code = code_match.group(1) if code_match else name[:4]
                prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

# ======== NIPPON ========
def parse_nippon():
    doc = docx.Document(os.path.join(BASE, "bang-gia-nippon-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 3: continue
            name = cells[0]
            if 'SẢN PHẨM' in name: continue
            if len(set(cells[2:4])) == 1 and cells[2] != '': continue
            spec = cells[2]
            price_raw = cells[3] if len(cells) > 3 else ''
            p = clean_price(price_raw)
            if p:
                code = f"NP{name[:6].upper()}"
                prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

# ======== SIKA ========
def parse_sika():
    doc = docx.Document(os.path.join(BASE, "bang-gia-sika-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            name = cells[0]
            if name.startswith('I.') or name.startswith('II.') or name.startswith('III.'): continue
            if name.startswith('1.') or name.startswith('2.') or name.startswith('3.'): continue
            if len(set(cells[2:5])) == 1: continue
            spec = cells[2]
            price_raw = cells[4] if len(cells) > 4 else cells[3]
            p = clean_price(price_raw)
            # Extract short code
            code_match = re.match(r'([\w\s\-]+?)(?:\s|$)', name)
            code = code_match.group(1).strip()[:15] if code_match else name[:12]
            if p:
                prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

# ======== GENERATE JS ========
results = {}
results['munich'] = parse_munich()
results['dulux'], results['maxilite'] = parse_dulux()
results['jotun'] = parse_jotun()
results['kova'] = parse_kova()
results['nano'] = parse_nano()
results['nippon'] = parse_nippon()
results['sika'] = parse_sika()

print("=== PRODUCT COUNTS (from approved DOCX) ===")
for k, v in results.items():
    print(f"{k}: {len(v)}")

# Generate BRANDS array
brand_map = {
    'munich': ('Munich', '⭐', '#e94560', 'Chống thấm & Sơn công nghiệp - Phân phối chính thức'),
    'nano': ('Nano House', '🏡', '#0fb9b1', 'Sơn & Chống thấm công nghệ Nano - Phân phối chính thức'),
    'sika': ('Sika', '🧪', '#6c5ce7', 'Vữa kỹ thuật & Chống thấm'),
    'dulux': ('Dulux', '🎨', '#0984e3', 'Sơn nội thất & ngoại thất'),
    'jotun': ('Jotun', '🖌️', '#f39c12', 'Sơn nội thất & ngoại thất cao cấp'),
    'kova': ('Kova', '🏺', '#e17055', 'Chống thấm & Sơn nước'),
    'nippon': ('Nippon', '🇯🇵', '#00b894', 'Sơn nội thất & ngoại thất'),
    'maxilite': ('Maxilite', '🔶', '#d4a017', 'Sơn kinh tế từ Dulux - AkzoNobel'),
    'mpe': ('MPE', '💡', '#2d3436', 'Thiết bị điện MPE'),
}

# Also load MPE data from the CAC file
import openpyxl
wb = openpyxl.load_workbook('/home/huu-minh/.openclaw/workspace/CAC_HANG_KHAC_TONG_HOP.xlsx')
ws = wb['MPE (Rạng Đông)']
mpe_count = 0
for row in ws.iter_rows(min_row=3, values_only=True):
    col1 = str(row[1]).strip() if row[1] else ''
    if col1 and col1 != 'Mã SP' and not col1.startswith('Nhóm'):
        mpe_count += 1
results['mpe'] = mpe_count
print(f"mpe: {mpe_count} (from CAC file)")

# Now generate the update commands
print("\n=== BRANDS ARRAY ===")
# The order: munich, nano, sika, dulux, jotun, kova, nippon, maxilite, mpe
order = ['munich', 'nano', 'sika', 'dulux', 'jotun', 'kova', 'nippon', 'maxilite', 'mpe']
for bid in order:
    v = results[bid]
    c = len(v) if isinstance(v, list) else v
    name, icon, color, desc = brand_map[bid]
    print(f"  {{ id:'{bid}', name:'{name}', icon:'{icon}', color:'{color}', desc:'{desc}', count:{c} }},")
