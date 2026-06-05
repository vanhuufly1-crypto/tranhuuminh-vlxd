#!/usr/bin/env python3
"""Parse DOCX báo giá chuẩn → update products.js + prices.js"""

import docx, re, json, os

BASE = "/media/huu-minh/DATA/tai-lieu/bảng giá các hãng/báo giá chuẩn"

def clean_price(s):
    s = s.replace('.','').replace(',','').replace('đ','').replace(' ','')
    m = re.findall(r'\d+', s)
    if m:
        return int(m[0])
    return None

def parse_munich():
    doc = docx.Document(os.path.join(BASE, "Bang-gia-Munich-TranHuuMinh.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            code = cells[0]
            name = cells[1]
            spec = cells[2]
            price_raw = cells[3]
            if not code or code == cells[1]: continue
            # price_raw may be "983.333đ / 2.950.000đ" → use first price
            first_price = price_raw.split('/')[0].strip()
            p = clean_price(first_price)
            prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

def parse_dulux():
    doc = docx.Document(os.path.join(BASE, "bang-gia-dulux-maxilite-22-05-2026.docx"))
    dulux_list = []
    maxilite_list = []
    in_maxilite = False
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            name = cells[0]
            if 'MAXILITE' in name.upper():
                in_maxilite = True
                continue
            if 'DULUX' in name.upper() or 'SƠN NGOẠI THẤT' in name.upper() or 'SƠN PHỦ NỘI THẤT' in name.upper():
                continue
            # Skip section headers
            if all(c == name for c in cells[2:] if c): continue
            code = cells[1] if len(cells) > 1 and cells[1] and 'Bao bì' not in cells[1] else name.split()[-1]
            spec = cells[2] if len(cells) > 2 else ''
            price_raw = cells[3] if len(cells) > 3 else ''
            p = clean_price(price_raw)
            item = {'code': code, 'name': name, 'spec': spec, 'price': p}
            if in_maxilite:
                maxilite_list.append(item)
            else:
                dulux_list.append(item)
    return dulux_list, maxilite_list

def parse_jotun():
    doc = docx.Document(os.path.join(BASE, "bang-gia-jotun-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            name = cells[0]
            if 'SẢN PHẨM' in name: continue
            if all(c == name for c in cells[2:] if c): continue
            code = name[:6]  # short code from name
            spec = cells[2] if len(cells) > 2 else ''
            price_raw = cells[3] if len(cells) > 3 else ''
            p = clean_price(price_raw)
            prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

def parse_kova():
    doc = docx.Document(os.path.join(BASE, "bang-gia-kova-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            code = cells[0]
            if not code: continue
            if code == cells[1] or 'SƠN NƯỚC' in code or 'MATÍT' in code: continue
            name = cells[1]
            spec = cells[2]
            price_raw = cells[3]
            p = clean_price(price_raw)
            prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

def parse_nano():
    doc = docx.Document(os.path.join(BASE, "bang-gia-nano-house-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 3: continue
            name = cells[0]
            if 'SẢN PHẨM' in name: continue
            if all(c == name for c in cells[2:] if c): continue
            # Extract code from name like "SUPER INTERIOR - NO1"
            code_match = re.search(r'[-–]\s*([\w\d]+)', name)
            code = code_match.group(1) if code_match else name[:4]
            spec = cells[2] if len(cells) > 2 else ''
            price_raw = cells[3] if len(cells) > 3 else cells[2]
            p = clean_price(price_raw)
            prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

def parse_nippon():
    doc = docx.Document(os.path.join(BASE, "bang-gia-nippon-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 3: continue
            name = cells[0]
            if 'SẢN PHẨM' in name: continue
            if all(c == name for c in cells[2:] if c): continue
            # Extract code
            code_match = re.search(r'NP\s+([\w\-]+)', name) or re.search(r'-\s*([\w\d]+)', name)
            code = f"NP-{code_match.group(1)[:8]}" if code_match else name[:6]
            spec = cells[2] if len(cells) > 2 else ''
            price_raw = cells[3] if len(cells) > 3 else ''
            p = clean_price(price_raw)
            prods.append({'code': code, 'name': name, 'spec': spec, 'price': p})
    return prods

def parse_sika():
    doc = docx.Document(os.path.join(BASE, "bang-gia-sika-22-05-2026.docx"))
    prods = []
    for table in doc.tables:
        for ri in range(2, len(table.rows)):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            if len(cells) < 4: continue
            name = cells[0]
            if name in ['I. SẢN PHẨM CHỐNG THẤM', 'II. SẢN PHẨM VỮA KỸ THUẬT']: continue
            if name.startswith('1.') or name.startswith('2.') or name.startswith('3.'): continue
            # Skip lines that are section headers
            if all(c == name for c in cells[2:] if c): continue
            # Use name as code+name since Sika products are well-known
            spec = cells[2]
            # Use giá bán sau CK (last column)
            price_raw = cells[4] if len(cells) > 4 else cells[3]
            p = clean_price(price_raw)
            prods.append({'code': name[:12], 'name': name, 'spec': spec, 'price': p})
    return prods

def fmt_price(n):
    if n is None: return "'Liên hệ'"
    return f"'{n:,}đ'.replace(',','.')"

# Run
results = {}
results['munich'] = parse_munich()
results['dulux'], results['maxilite'] = parse_dulux()
results['jotun'] = parse_jotun()
results['kova'] = parse_kova()
results['nano'] = parse_nano()
results['nippon'] = parse_nippon()
results['sika'] = parse_sika()

for k, v in results.items():
    count = len(v) if isinstance(v, list) else 0
    print(f"{k}: {count}")
    if count > 0:
        print(f"  first: {v[0]}")
        print(f"  last: {v[-1]}")
